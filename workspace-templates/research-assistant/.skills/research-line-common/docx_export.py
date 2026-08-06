#!/usr/bin/env python3
"""Export research-line JSON outputs to teacher-facing DOCX review documents."""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SUPPORTED_SKILLS = {
    "research-topic-generation-skill",
    "literature-reading-skill",
    "paper-writing-skill",
    "project-proposal-skill",
}


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def slug(value: str) -> str:
    text = re.sub(r"[^A-Za-z0-9_.-]+", "-", value).strip("-")
    return text[:80] or "research-output"


def text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return json.dumps(value, ensure_ascii=False)


def fact_value(value: Any) -> str:
    if isinstance(value, list):
        return "；".join(text(item) for item in value)
    return text(value)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.add_run(subtitle)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "E8EEF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_bullets(doc: Document, values: list[Any]) -> None:
    for value in values:
        doc.add_paragraph(text(value), style="List Bullet")


def add_data_source_section(doc: Document, data: dict[str, Any]) -> None:
    report = data.get("dataSourceReport", {})
    sources = report.get("dataSources", []) if isinstance(report, dict) else []
    if not sources:
        return
    doc.add_heading("数据源与使用边界", level=1)
    rows = [
        [
            item.get("sourceName", ""),
            item.get("sourceType", ""),
            item.get("dataType", ""),
            str(item.get("recordCount", "")),
            item.get("authorizationStatus", ""),
            "；".join(item.get("limitations", [])),
        ]
        for item in sources
        if isinstance(item, dict)
    ]
    add_table(doc, ["数据源", "类型", "数据", "数量", "授权状态", "限制"], rows)
    limitations = report.get("overallLimitations", [])
    if limitations:
        doc.add_paragraph("总体限制：")
        add_bullets(doc, limitations)


def export_research_topic(data: dict[str, Any], output_dir: Path) -> list[Path]:
    request_id = slug(data.get("requestId", "research-topic"))
    output_path = output_dir / f"{request_id}-research-topic-report.docx"
    result = data.get("result", {})
    doc = Document()
    configure_document(doc)
    add_title(doc, "研究选题生成报告", f"请求 ID：{data.get('requestId', '')}")
    doc.add_paragraph(f"质量状态：{data.get('qualityReport', {}).get('status', '')}")
    doc.add_paragraph(data.get("summary", ""))

    add_data_source_section(doc, data)

    topics = result.get("topicCandidates", [])
    if topics:
        doc.add_heading("候选选题", level=1)
        rows = [
            [
                item.get("topicId", ""),
                item.get("topicTitle", ""),
                item.get("topicType", ""),
                "、".join(item.get("keywords", [])),
                str(item.get("feasibility", {}).get("score", "")),
                item.get("differentiation", {}).get("riskLevel", ""),
            ]
            for item in topics
        ]
        add_table(doc, ["ID", "题目", "类型", "关键词", "可行性", "重复风险"], rows)
        for item in topics:
            doc.add_heading(item.get("topicTitle", ""), level=2)
            doc.add_paragraph(f"研究问题：{item.get('researchQuestion', '')}")
            existing_basis = item.get("existingBasis", [])
            if existing_basis:
                doc.add_paragraph("已有基础：")
                add_bullets(doc, [basis.get("basis", basis) if isinstance(basis, dict) else basis for basis in existing_basis])
            basis_gap = item.get("basisGap", {})
            if basis_gap:
                doc.add_paragraph(f"当前基础：{'；'.join(basis_gap.get('currentBasis', []))}")
                doc.add_paragraph(f"目标要求：{basis_gap.get('targetRequirement', '')}")
                if basis_gap.get("gaps"):
                    doc.add_paragraph("缺口：")
                    add_bullets(doc, basis_gap.get("gaps", []))
                if basis_gap.get("upgradePath"):
                    doc.add_paragraph("升级路径：")
                    add_bullets(doc, basis_gap.get("upgradePath", []))

    clusters = result.get("materialClusters", [])
    if clusters:
        doc.add_heading("材料主题聚类", level=1)
        rows = [
            [
                item.get("clusterId", ""),
                item.get("clusterTitle", ""),
                "、".join(item.get("materialIds", [])),
                "、".join(item.get("coreSignals", [])),
                item.get("currentStage", ""),
            ]
            for item in clusters
        ]
        add_table(doc, ["聚类", "主题", "材料", "信号", "阶段"], rows)

    trajectory = result.get("researchTrajectory", {})
    if trajectory:
        doc.add_heading("研究轨迹", level=1)
        doc.add_paragraph(f"阶段：{trajectory.get('stage', '')}")
        doc.add_paragraph(trajectory.get("trajectorySummary", ""))
        if trajectory.get("futureDeepeningPath"):
            doc.add_paragraph("后续深化路径：")
            add_bullets(doc, [step.get("action", step) if isinstance(step, dict) else step for step in trajectory.get("futureDeepeningPath", [])])

    next_actions = data.get("nextActions", [])
    if next_actions:
        doc.add_heading("下一步", level=1)
        add_bullets(doc, next_actions)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return [output_path]


def export_literature_reading(data: dict[str, Any], output_dir: Path) -> list[Path]:
    request_id = slug(data.get("requestId", "literature-reading"))
    output_path = output_dir / f"{request_id}-literature-reading-report.docx"
    result = data.get("result", {})
    doc = Document()
    configure_document(doc)
    add_title(doc, "文献阅读助手报告", f"请求 ID：{data.get('requestId', '')}")
    doc.add_paragraph(f"质量状态：{data.get('qualityReport', {}).get('status', '')}")
    doc.add_paragraph(data.get("summary", ""))

    add_data_source_section(doc, data)

    search_report = result.get("corpusSearchReport", {})
    if search_report:
        doc.add_heading("检索报告", level=1)
        query = search_report.get("query", {})
        doc.add_paragraph(f"研究主题：{query.get('researchTopic', '')}")
        doc.add_paragraph(f"关键词：{'、'.join(query.get('keywords', []))}")
        doc.add_paragraph(f"候选数：{search_report.get('candidateCount', 0)}；返回数：{search_report.get('returnedCount', 0)}")

    records = result.get("literatureRecords", [])
    if records:
        doc.add_heading("推荐阅读列表", level=1)
        rows = [
            [
                item.get("paperId", ""),
                item.get("title", ""),
                "、".join(item.get("authors", [])),
                str(item.get("year", "")),
                item.get("textAvailability", ""),
                item.get("sourceStatus", ""),
            ]
            for item in records
        ]
        add_table(doc, ["paperId", "题名", "作者", "年份", "文本", "来源状态"], rows)

    quick_cards = result.get("quickReadCards", [])
    if quick_cards:
        doc.add_heading("速读卡", level=1)
        rows = [
            [
                item.get("cardId", ""),
                item.get("paperId", ""),
                item.get("readingDecision", ""),
                item.get("topicRelevance", ""),
                item.get("findings", ""),
                item.get("limitations", ""),
            ]
            for item in quick_cards
        ]
        add_table(doc, ["卡片", "paperId", "阅读决策", "相关度", "主要信息", "限制"], rows)

    deep_cards = result.get("deepReadCards", [])
    if deep_cards:
        doc.add_heading("精读卡", level=1)
        for item in deep_cards:
            doc.add_heading(item.get("paperId", ""), level=2)
            doc.add_paragraph(f"问题：{item.get('researchProblem', '')}")
            doc.add_paragraph(f"方法：{item.get('method', '')}")
            doc.add_paragraph("发现：")
            add_bullets(doc, item.get("findings", []))
            doc.add_paragraph("限制：")
            add_bullets(doc, item.get("limitations", []))

    comparison = result.get("comparisonMatrix", [])
    if comparison:
        doc.add_heading("横向比较", level=1)
        for matrix in comparison:
            doc.add_heading(matrix.get("topic", ""), level=2)
            rows = [
                [
                    row.get("paperId", ""),
                    row.get("problem", ""),
                    row.get("method", ""),
                    row.get("finding", ""),
                    row.get("limitation", ""),
                ]
                for row in matrix.get("rows", [])
            ]
            add_table(doc, ["paperId", "问题", "方法", "发现", "限制"], rows)

    evidence_cards = result.get("evidenceCards", [])
    if evidence_cards:
        doc.add_heading("可复用证据卡", level=1)
        rows = [
            [
                item.get("cardId", ""),
                item.get("paperId", ""),
                item.get("supportType", ""),
                item.get("evidenceLevel", ""),
                item.get("claim", ""),
                "；".join(item.get("limits", [])),
            ]
            for item in evidence_cards
        ]
        add_table(doc, ["证据卡", "paperId", "支撑类型", "证据级别", "主张", "边界"], rows)

    next_actions = data.get("nextActions", [])
    if next_actions:
        doc.add_heading("下一步", level=1)
        add_bullets(doc, next_actions)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return [output_path]


def export_paper(data: dict[str, Any], output_dir: Path) -> list[Path]:
    request_id = slug(data.get("requestId", "paper-writing"))
    output_path = output_dir / f"{request_id}-paper-writing-review.docx"
    result = data.get("result", {})
    doc = Document()
    configure_document(doc)
    add_title(doc, "论文写作助手查源与润色报告", f"请求 ID：{data.get('requestId', '')}")
    doc.add_paragraph(f"质量状态：{data.get('qualityReport', {}).get('status', '')}")

    trace = (result.get("sourceTraceResults") or [{}])[0]
    doc.add_heading("查源结论", level=1)
    add_table(
        doc,
        ["查询文本", "判定"],
        [[trace.get("queryText", ""), trace.get("decision", "")]],
    )

    candidates = trace.get("candidates", [])
    if candidates:
        doc.add_heading("候选文献", level=1)
        rows = [
            [
                item.get("paperId", ""),
                item.get("supportStatus", ""),
                item.get("evidenceLevel", ""),
                item.get("citation", ""),
            ]
            for item in candidates
        ]
        add_table(doc, ["paperId", "支撑状态", "证据级别", "引用格式"], rows)

    claim_checks = result.get("claimChecks", [])
    if claim_checks:
        doc.add_heading("论点检查", level=1)
        rows = [
            [
                item.get("claimId", ""),
                item.get("claimText", ""),
                item.get("status", ""),
                "；".join(item.get("riskNotes", [])),
                item.get("recommendedRewrite", ""),
            ]
            for item in claim_checks
        ]
        add_table(doc, ["论点", "文本", "状态", "风险", "建议"], rows)

    structure = result.get("structureDiagnosis", {})
    if structure:
        doc.add_heading("结构诊断", level=1)
        for item in structure.get("sectionCoverage", []):
            weak = "、".join(item.get("weakElements", []) or item.get("missingElements", [])) or "暂无"
            doc.add_paragraph(f"{item.get('label')}：{item.get('status')}；关注：{weak}", style="List Bullet")
        missing = [item.get("label") for item in structure.get("abstractChecklist", []) if item.get("status") == "missing"]
        if missing:
            doc.add_paragraph(f"摘要缺失要素：{'、'.join(missing)}")

    suggestions = result.get("revisionSuggestions", [])
    if suggestions:
        doc.add_heading("保守润色建议", level=1)
        rows = [
            [
                item.get("suggestionId", ""),
                item.get("editType", ""),
                item.get("originalText", ""),
                item.get("revisedText", ""),
                "是" if item.get("needsEvidence") else "否",
            ]
            for item in suggestions
        ]
        add_table(doc, ["ID", "类型", "原文", "建议", "需补证据"], rows)

    insertions = result.get("insertionSuggestions", [])
    if insertions:
        doc.add_heading("待教师确认插入建议", level=1)
        rows = [
            [
                item.get("insertionId", ""),
                item.get("claimId", ""),
                item.get("status", ""),
                item.get("formattedCitation", ""),
            ]
            for item in insertions
        ]
        add_table(doc, ["ID", "论点", "状态", "参考文献"], rows)

    next_actions = data.get("nextActions", [])
    if next_actions:
        doc.add_heading("下一步", level=1)
        add_bullets(doc, next_actions)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return [output_path]


def project_title(fact_table: dict[str, Any]) -> str:
    for fact in fact_table.get("facts", []):
        if fact.get("field") == "project.title":
            return fact_value(fact.get("value"))
    return "项目文档"


def export_project_document(data: dict[str, Any], document_data: dict[str, Any], output_dir: Path) -> Path:
    request_id = slug(data.get("requestId", "project-proposal"))
    document_type = slug(document_data.get("documentType", "document"))
    output_path = output_dir / f"{request_id}-{document_type}.docx"
    result = data.get("result", {})
    fact_table = result.get("projectFactTable", {})

    doc = Document()
    configure_document(doc)
    add_title(doc, document_data.get("title", "项目文档"), f"文档类型：{document_data.get('documentType', '')}")
    doc.add_paragraph(f"项目题目：{project_title(fact_table)}")
    doc.add_paragraph(f"质量状态：{data.get('qualityReport', {}).get('status', '')}")

    doc.add_heading("正文框架", level=1)
    for section in document_data.get("sections", []):
        doc.add_heading(section.get("title", ""), level=2)
        doc.add_paragraph(section.get("content", ""))
        doc.add_paragraph(f"状态：{section.get('status', '')}")
        doc.add_paragraph(f"事实引用：{', '.join(section.get('factRefs', [])) or '无'}")

    doc.add_heading("项目事实表", level=1)
    rows = [
        [
            fact.get("factId", ""),
            fact.get("field", ""),
            fact_value(fact.get("value")),
            ", ".join(fact.get("sourceRefs", [])),
            fact.get("status", ""),
        ]
        for fact in fact_table.get("facts", [])
    ]
    add_table(doc, ["factId", "字段", "值", "来源", "状态"], rows)

    if fact_table.get("missingFields"):
        doc.add_heading("缺失字段", level=1)
        for item in fact_table.get("missingFields", []):
            if isinstance(item, dict):
                doc.add_paragraph(f"{item.get('field')}：{item.get('reason', '')}", style="List Bullet")
            else:
                doc.add_paragraph(text(item), style="List Bullet")

    if fact_table.get("conflicts"):
        doc.add_heading("冲突字段", level=1)
        for item in fact_table.get("conflicts", []):
            doc.add_paragraph(f"{item.get('field')}：{item.get('resolution', '')}", style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def export_project(data: dict[str, Any], output_dir: Path) -> list[Path]:
    result = data.get("result", {})
    document = result.get("documentDraft", {})
    document_set = result.get("documentSet", {})
    documents = []
    if isinstance(document_set, dict) and isinstance(document_set.get("documents"), list):
        documents = document_set["documents"]
    elif isinstance(document, dict) and document:
        documents = [document]
    return [export_project_document(data, item, output_dir) for item in documents]


def export_one(path: Path, output_dir: Path) -> list[Path]:
    data = read_json(path)
    skill_id = data.get("skillId")
    if skill_id not in SUPPORTED_SKILLS:
        raise ValueError(f"暂不支持导出 skillId={skill_id} 的 DOCX。")
    if skill_id == "research-topic-generation-skill":
        return export_research_topic(data, output_dir)
    if skill_id == "literature-reading-skill":
        return export_literature_reading(data, output_dir)
    if skill_id == "paper-writing-skill":
        return export_paper(data, output_dir)
    return export_project(data, output_dir)


def main() -> int:
    parser = argparse.ArgumentParser(description="将科研线 JSON 输出导出为 DOCX 交付物。")
    parser.add_argument("output_json", nargs="+", help="科研线 JSON 输出")
    parser.add_argument("--output-dir", required=True, help="DOCX 输出目录")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    exported: list[Path] = []
    for item in args.output_json:
        exported.extend(export_one(Path(item), output_dir))
    payload = {
        "status": "passed",
        "outputDir": str(output_dir),
        "files": [{"path": str(path), "bytes": path.stat().st_size} for path in exported],
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
